import os
import base64
from openai import OpenAI
from pdf2docx import Converter
from docx import Document
from spellchecker import SpellChecker
from flask import Blueprint, request, jsonify, send_file
from flask_jwt_extended import jwt_required, get_jwt_identity
from app.models.user_model import User
from app.models.resume_model import Resume
from app import db
import io
from flask_cors import cross_origin
import win32com.client
import jwt


resume_routes = Blueprint('resume_routes', __name__)
UPLOAD_FOLDER = os.path.join(os.getcwd(), 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# ✅ Get User Resumes
@resume_routes.route('/list', methods=['GET'])
@jwt_required()
def list_resumes():
    user_id = get_jwt_identity()
    resumes = Resume.query.filter_by(user_id=user_id).all()
    resume_list = [
        {
            "id": r.id,
            "title": r.title,
            "file_url": f"http://localhost:5001/resume/download/{r.id}"
        }
        for r in resumes
    ]

    return jsonify(resume_list)



@resume_routes.route("/user-profile", methods=["GET"])
@jwt_required()
def get_user_profile():
    user_id = get_jwt_identity()
    user = db.session.get(User, user_id)

    if not user:
        return jsonify({"error": "User not found"}), 404

    # 🖼️ Convert binary profile pic to Base64 string
    profile_pic_base64 = (
        base64.b64encode(user.profile_pic).decode('utf-8') 
        if user.profile_pic 
        else None
    )

    return jsonify({
        "name": user.name,
        "profile_pic": profile_pic_base64
    })

# ✅ Upload Resume with Binary File Storage
@resume_routes.route("/upload", methods=["POST"])
@jwt_required()
def upload_resume():
    user_id = get_jwt_identity()
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"error": "User not found"}), 404

    file = request.files.get("resume_file")
    title = request.form.get("title")

    if not file or not title:
        return jsonify({"error": "Both file and title are required"}), 400

    file_data = file.read()

    new_resume = Resume(user_id=user.id, title=title, file_data=file_data)
    db.session.add(new_resume)
    db.session.commit()

    return jsonify({"message": "Resume uploaded successfully!"})


# ✅ Adjust user-resumes to include the PDF URL
@resume_routes.route("/user-resumes", methods=["GET"])
@jwt_required()
def get_user_resumes():
    user_id = get_jwt_identity()
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"error": "User not found"}), 404

    resumes = [
        {
            "id": r.id,
            "title": r.title,
            "file_url": f"http://localhost:5001/resume/download/{r.id}"
        }
        for r in user.resumes
    ]

    return jsonify(resumes)


# ✅ Upload Profile Picture (Directly into DB as Binary Data)
@resume_routes.route("/upload-profile-pic", methods=["POST"])
@jwt_required()
def upload_profile_pic():
    user_id = get_jwt_identity()
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"error": "User not found"}), 404

    file = request.files.get("profile_pic")
    if not file:
        return jsonify({"error": "Profile picture is required"}), 400

    file_data = file.read()
    user.profile_pic = file_data  # Save binary data in DB
    db.session.commit()

    profile_pic_base64 = base64.b64encode(file_data).decode('utf-8')

    return jsonify({"message": "Profile picture uploaded successfully", "profile_pic": profile_pic_base64})



import logging
logging.basicConfig(level=logging.DEBUG)
# 🛠️ Endpoint to serve PDF files
@resume_routes.route("/download/<int:resume_id>", methods=["GET"])
@jwt_required()
def download_resume(resume_id):
    # Verify user identity
    user_id = get_jwt_identity()

    # Query the resume
    resume = Resume.query.get(resume_id)
    if not resume:
        return jsonify({"error": "Resume not found"}), 404

    # Ensure the resume belongs to the authenticated user
    if resume.user_id != int(user_id):
        return jsonify({"error": "Unauthorized"}), 403

    # Serve the PDF from binary data
    return send_file(
        io.BytesIO(resume.file_data),
        mimetype="application/pdf",
        as_attachment=False,
        download_name=resume.title
    )


@resume_routes.route("/enhance/<int:resume_id>", methods=["POST"])
@cross_origin(supports_credentials=True)
@jwt_required()
def enhance_resume(resume_id):
    logging.debug(f"📌 Received enhancement request for resume ID: {resume_id}")

    try:
        resume = db.session.get(Resume, resume_id)
        if not resume:
            logging.error("🚨 Resume not found.")
            return jsonify({"error": "Resume not found"}), 404

        # ✅ Save PDF Temporarily
        temp_pdf_path = os.path.join(UPLOAD_FOLDER, f"resume_{resume_id}.pdf")
        if os.path.exists(temp_pdf_path):
            os.remove(temp_pdf_path)  # Remove existing file

        with open(temp_pdf_path, "wb") as f:
            f.write(resume.file_data)

        # ✅ Convert PDF to DOCX
        docx_path = temp_pdf_path.replace(".pdf", ".docx")
        cv = Converter(temp_pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()

        # ✅ Extract text from DOCX
        doc = Document(docx_path)
        original_text = "\n".join([p.text for p in doc.paragraphs])
        if not original_text.strip():
            return jsonify({"error": "No text found in resume"}), 400

        # ✅ Construct AI Prompt
        prompt = (
            "Enhance the following resume text. Correct spelling mistakes, improve grammar, and "
            "make it ATS-friendly while preserving its original formatting. Provide a structured response "
            "in the exact format below:\n\n"
            
            "**Enhanced Text:**\n"
            "<Provide the fully rewritten resume text here>\n\n"
            
            "**Improvements:**\n"
            "- <Bullet point of improvement>\n"
            "- <Bullet point of improvement>\n\n"
            
            "**Changes (Before → After):**\n"
            "<Original sentence> → <Enhanced sentence>\n"
            "<Original sentence> → <Enhanced sentence>\n\n"
            
            "Resume Content:\n"
            f"{original_text}"
        )

        # ✅ Send Request to OpenAI
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "system", "content": "You are a professional resume enhancer."},
                      {"role": "user", "content": prompt}],
            max_tokens=2048
        )

        ai_response = response.choices[0].message.content.strip()
        logging.debug(f"📌 AI Response Preview: {ai_response[:100]}...")

        # ✅ Parse AI Response (Extract Sections)
        sections = {"enhanced_text": "", "improvements": [], "changes": []}

        if "**Enhanced Text:**" in ai_response:
            parts = ai_response.split("**Enhanced Text:**")[1]

            if "**Improvements:**" in parts:
                sections["enhanced_text"], parts = parts.split("**Improvements:**", 1)
            if "**Changes (Before → After):**" in parts:
                sections["improvements"], sections["changes"] = parts.split("**Changes (Before → After):**", 1)

            sections["enhanced_text"] = sections["enhanced_text"].strip()
            sections["improvements"] = [i.strip("- ") for i in sections["improvements"].strip().split("\n") if i]
            sections["changes"] = [{"before": c.split("→")[0].strip(), "after": c.split("→")[1].strip()} 
                                   for c in sections["changes"].strip().split("\n") if "→" in c]
        else:
            return jsonify({"error": "AI response missing required sections"}), 500

        # ✅ Insert Enhanced Text Back into DOCX (Only Resume Content)
        doc.paragraphs.clear()
        for line in sections["enhanced_text"].split("\n"):
            doc.add_paragraph(line)
        doc.save(docx_path)

        # ✅ Convert DOCX Back to PDF
        final_pdf_path = os.path.join(UPLOAD_FOLDER, f"resume_{resume_id}-enhanced.pdf")

        import pythoncom

        def convert_docx_to_pdf(docx_path, pdf_path):
            pythoncom.CoInitialize()  # ✅ Initialize COM

            try:
                word = win32com.client.Dispatch("Word.Application")
                doc = word.Documents.Open(docx_path)
                doc.SaveAs(pdf_path, FileFormat=17)  # 17 = PDF format
                doc.Close()
                word.Quit()
            except Exception as e:
                print(f"Error during DOCX to PDF conversion: {e}")
            finally:
                pythoncom.CoUninitialize()  # ✅ Uninitialize COM properly

        convert_docx_to_pdf(docx_path, final_pdf_path)

        # ✅ Handle Windows Auto-Adding `_1` to Filenames
        enhanced_filename = f"resume_{resume_id}-enhanced.pdf"
        if not os.path.exists(final_pdf_path):
            alternative_filename = os.path.join(UPLOAD_FOLDER, f"resume_{resume_id}-enhanced_1.pdf")
            if os.path.exists(alternative_filename):
                final_pdf_path = alternative_filename
                enhanced_filename = f"resume_{resume_id}-enhanced_1.pdf"

        # ✅ Read Enhanced PDF into Binary
        with open(final_pdf_path, "rb") as f:
            enhanced_pdf_data = f.read()

        # ✅ Save to database
        resume.enhanced_text = sections["enhanced_text"]
        resume.improvements = ",".join(sections["improvements"])
        resume.changes = sections["changes"]
        resume.enhanced_filename = enhanced_filename
        resume.enhanced_file_data = enhanced_pdf_data

        db.session.commit()

        return jsonify({
            "message": "Resume enhanced successfully!",
            "improvements": sections["improvements"],
            "changes": sections["changes"],
            "download_url": f"http://localhost:5001/resume/download/{enhanced_filename}"
        })

    except Exception as e:
        return jsonify({"error": f"Unexpected error: {str(e)}"}), 500


@resume_routes.route("/download/<string:filename>", methods=["GET"])
def download_enhanced_resume(filename):
    """Serves the enhanced resume PDF with token authentication."""

    token = request.args.get("token")
    if not token:
        logging.error("Missing token in download request.")
        return jsonify({"error": "Missing token"}), 401

    # Verify JWT Token
    try:
        decoded_token = jwt.decode(token, options={"verify_signature": False})
        user_id = decoded_token.get("sub")
        logging.info(f"User {user_id} attempting to download {filename}")
    except Exception as e:
        logging.error(f"Invalid token: {str(e)}")
        return jsonify({"error": f"Invalid token: {str(e)}"}), 401

    # Ensure the filename includes .pdf
    if not filename.endswith(".pdf"):
        filename += ".pdf"

    # Fetch resume from database
    resume = db.session.query(Resume).filter(Resume.enhanced_filename == filename).first()
    if not resume:
        logging.error(f"🚨 Resume {filename} not found in database.")
        return jsonify({"error": "Unauthorized or file not found"}), 403

    if resume.user_id != int(user_id):
        logging.error(f"🚨 User {user_id} is not authorized to access {filename}")
        return jsonify({"error": "Unauthorized access"}), 403

    # ✅ Handle Windows Auto-Adding `_1` to Filenames
    file_path = os.path.join(UPLOAD_FOLDER, filename)
    if not os.path.exists(file_path):
        alternative_filename = os.path.join(UPLOAD_FOLDER, filename.replace(".pdf", "_1.pdf"))
        if os.path.exists(alternative_filename):
            file_path = alternative_filename
            logging.info(f"✅ Found alternative filename: {file_path}")

    if not os.path.exists(file_path):
        logging.error(f"🚨 File not found: {file_path}")
        return jsonify({"error": "File not found"}), 404

    logging.info(f"✅ Serving file: {file_path} for user {user_id}")
    return send_file(file_path, mimetype="application/pdf", as_attachment=True)

